from docx import Document
import os
import logging

class ExportService:
    def export_to_docx(self, content, filename):
        # Get the absolute path to the project root
        project_root = os.path.abspath(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
        exports_dir = os.path.join(project_root, 'exports')
        
        # Create exports directory if it doesn't exist
        os.makedirs(exports_dir, exist_ok=True)
        
        # Create absolute file path
        file_path = os.path.join(exports_dir, f"{filename}.docx")
        
        # Log the path for debugging
        logging.info(f"Saving DOCX to: {file_path}")
        
        # Create a new Document
        doc = Document()
        
        # Add content to the document
        doc.add_heading('User Story', 0)
        
        doc.add_heading('Role', level=1)
        doc.add_paragraph(content['role'])
        
        doc.add_heading('Goal', level=1)
        doc.add_paragraph(content['goal'])
        
        doc.add_heading('Benefit', level=1)
        doc.add_paragraph(content['benefit'])
        
        doc.add_heading('Acceptance Criteria', level=1)
        for criterion in content['acceptance_criteria']:
            doc.add_paragraph(criterion, style='List Bullet')
        
        # Save the document
        doc.save(file_path)
        
        print(f"File saved to: {file_path}")  # Console output for immediate feedback
        return file_path
